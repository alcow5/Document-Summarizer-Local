"""
Document Processing Service
Handles text extraction from PDF and DOCX files with chunking for large documents
"""

import PyPDF2
import docx
import io
import re
from typing import List, Dict, Any, Optional
from loguru import logger
import tiktoken
from pathlib import Path


class DocumentProcessor:
    """Service for processing and extracting text from documents"""
    
    def __init__(self, config: Dict[str, Any]):
        self.config = config
        self.chunk_size = config["documents"]["chunk_size"]
        self.overlap_size = config["documents"]["overlap_size"]
        self.max_pages = config["documents"]["max_pages"]
        self.supported_formats = config["documents"]["supported_formats"]
        
        # Initialize tokenizer for accurate chunking
        try:
            self.tokenizer = tiktoken.get_encoding("cl100k_base")
        except Exception as e:
            logger.warning(f"Failed to load tokenizer: {e}")
            self.tokenizer = None
    
    async def extract_text(self, file_content: bytes, file_extension: str) -> List[str]:
        """
        Extract text from document and return chunks
        
        Args:
            file_content: Document file content as bytes
            file_extension: File extension (pdf, docx)
            
        Returns:
            List of text chunks
        """
        try:
            if file_extension.lower() == "pdf":
                text = await self._extract_pdf_text(file_content)
            elif file_extension.lower() == "docx":
                text = await self._extract_docx_text(file_content)
            else:
                raise ValueError(f"Unsupported file format: {file_extension}")
            
            # Clean and normalize text
            text = self._clean_text(text)
            
            # Split into chunks
            chunks = self._create_chunks(text)
            
            logger.info(f"Extracted {len(text)} characters, created {len(chunks)} chunks")
            return chunks
            
        except Exception as e:
            logger.error(f"Error extracting text from {file_extension} file: {e}")
            raise
    
    async def _extract_pdf_text(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Check page count
            page_count = len(pdf_reader.pages)
            if page_count > self.max_pages:
                logger.warning(f"PDF has {page_count} pages, limiting to {self.max_pages}")
                page_count = self.max_pages
            
            text_parts = []
            
            for page_num in range(page_count):
                try:
                    page = pdf_reader.pages[page_num]
                    page_text = page.extract_text()
                    
                    if page_text.strip():
                        text_parts.append(page_text)
                        logger.debug(f"Extracted {len(page_text)} characters from page {page_num + 1}")
                    
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            full_text = "\n\n".join(text_parts)
            logger.info(f"Successfully extracted text from PDF: {len(full_text)} characters from {page_count} pages")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error processing PDF: {e}")
            raise ValueError(f"Failed to process PDF file: {str(e)}")
    
    async def _extract_docx_text(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)
            
            text_parts = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(" | ".join(row_text))
            
            full_text = "\n\n".join(text_parts)
            logger.info(f"Successfully extracted text from DOCX: {len(full_text)} characters")
            
            return full_text
            
        except Exception as e:
            logger.error(f"Error processing DOCX: {e}")
            raise ValueError(f"Failed to process DOCX file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        
        # Remove multiple consecutive newlines
        text = re.sub(r'\n\s*\n\s*\n+', '\n\n', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\xff]', '', text)
        
        # Strip whitespace
        text = text.strip()
        
        return text
    
    def _create_chunks(self, text: str) -> List[str]:
        """
        Split text into chunks with overlap for processing by LLM
        
        Args:
            text: Full document text
            
        Returns:
            List of text chunks
        """
        if not text:
            return []
        
        # If text is small enough, return as single chunk
        if self.tokenizer:
            token_count = len(self.tokenizer.encode(text))
            if token_count <= self.chunk_size:
                return [text]
        else:
            # Fallback: estimate tokens as ~4 characters per token
            estimated_tokens = len(text) // 4
            if estimated_tokens <= self.chunk_size:
                return [text]
        
        chunks = []
        
        # Split text into sentences for better chunking
        sentences = self._split_into_sentences(text)
        
        current_chunk = ""
        current_tokens = 0
        
        for sentence in sentences:
            sentence_tokens = self._count_tokens(sentence)
            
            # If adding this sentence would exceed chunk size
            if current_tokens + sentence_tokens > self.chunk_size and current_chunk:
                chunks.append(current_chunk.strip())
                
                # Start new chunk with overlap
                overlap_text = self._get_overlap_text(current_chunk)
                current_chunk = overlap_text + " " + sentence
                current_tokens = self._count_tokens(current_chunk)
            else:
                current_chunk += " " + sentence if current_chunk else sentence
                current_tokens += sentence_tokens
        
        # Add the last chunk
        if current_chunk.strip():
            chunks.append(current_chunk.strip())
        
        logger.info(f"Created {len(chunks)} chunks from text")
        return chunks
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """Split text into sentences for better chunking"""
        # Simple sentence splitting - can be enhanced with NLTK if needed
        sentences = re.split(r'(?<=[.!?])\s+', text)
        return [s.strip() for s in sentences if s.strip()]
    
    def _count_tokens(self, text: str) -> int:
        """Count tokens in text"""
        if self.tokenizer:
            return len(self.tokenizer.encode(text))
        else:
            # Fallback estimation
            return len(text) // 4
    
    def _get_overlap_text(self, chunk: str) -> str:
        """Get overlap text from the end of current chunk"""
        if not chunk:
            return ""
        
        # Get sentences from the end of the chunk for overlap
        sentences = self._split_into_sentences(chunk)
        if not sentences:
            return ""
        
        overlap_text = ""
        overlap_tokens = 0
        
        # Add sentences from the end until we reach overlap size
        for sentence in reversed(sentences):
            sentence_tokens = self._count_tokens(sentence)
            if overlap_tokens + sentence_tokens <= self.overlap_size:
                overlap_text = sentence + " " + overlap_text if overlap_text else sentence
                overlap_tokens += sentence_tokens
            else:
                break
        
        return overlap_text.strip()
    
    def get_document_metadata(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Extract metadata from document
        
        Args:
            file_content: Document file content
            filename: Original filename
            
        Returns:
            Dictionary with document metadata
        """
        file_extension = Path(filename).suffix.lower().lstrip('.')
        
        metadata = {
            "filename": filename,
            "file_size": len(file_content),
            "file_extension": file_extension,
            "page_count": None,
            "word_count": None,
            "language": None
        }
        
        try:
            # Extract text to get word count
            if file_extension == "pdf":
                text = self._extract_pdf_text_sync(file_content)
                # Get page count for PDF
                pdf_file = io.BytesIO(file_content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                metadata["page_count"] = len(pdf_reader.pages)
                
            elif file_extension == "docx":
                text = self._extract_docx_text_sync(file_content)
                # For DOCX, estimate pages based on word count
                word_count = len(text.split())
                metadata["page_count"] = max(1, word_count // 250)  # ~250 words per page
            
            # Calculate word count
            if text:
                metadata["word_count"] = len(text.split())
                
                # Simple language detection (can be enhanced)
                metadata["language"] = self._detect_language(text)
            
        except Exception as e:
            logger.warning(f"Failed to extract metadata: {e}")
        
        return metadata
    
    def _extract_pdf_text_sync(self, file_content: bytes) -> str:
        """Synchronous version of PDF text extraction for metadata"""
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        for page in pdf_reader.pages[:5]:  # Only first 5 pages for metadata
            text_parts.append(page.extract_text())
        
        return "\n".join(text_parts)
    
    def _extract_docx_text_sync(self, file_content: bytes) -> str:
        """Synchronous version of DOCX text extraction for metadata"""
        docx_file = io.BytesIO(file_content)
        doc = docx.Document(docx_file)
        
        text_parts = []
        for paragraph in doc.paragraphs[:50]:  # Only first 50 paragraphs for metadata
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        return "\n".join(text_parts)
    
    def _detect_language(self, text: str) -> str:
        """Simple language detection based on common words"""
        # This is a very basic implementation
        # In production, consider using langdetect or similar libraries
        
        text_lower = text.lower()
        
        # English indicators
        english_words = ['the', 'and', 'of', 'to', 'a', 'in', 'is', 'it', 'you', 'that']
        english_count = sum(1 for word in english_words if word in text_lower)
        
        # Spanish indicators
        spanish_words = ['el', 'de', 'que', 'y', 'a', 'en', 'un', 'es', 'se', 'no']
        spanish_count = sum(1 for word in spanish_words if word in text_lower)
        
        # French indicators
        french_words = ['le', 'de', 'et', 'à', 'un', 'il', 'être', 'et', 'en', 'avoir']
        french_count = sum(1 for word in french_words if word in text_lower)
        
        if english_count >= spanish_count and english_count >= french_count:
            return "en"
        elif spanish_count >= french_count:
            return "es"
        elif french_count > 0:
            return "fr"
        else:
            return "unknown" 